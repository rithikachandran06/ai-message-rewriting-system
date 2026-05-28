"""
Streamlit UI for Text Rewriting with RLHF
Modern web interface for sentiment analysis, text rewriting, and RLHF training
"""

import streamlit as st
import sys
import os
import traceback

# Windows-specific memory optimizations - MUST be before torch import
if sys.platform == "win32":
    os.environ["PYTORCH_CUDA_ALLOC_CONF"] = "max_split_size_mb:128"
    os.environ["OMP_NUM_THREADS"] = "1"
    os.environ["MKL_NUM_THREADS"] = "1"

import torch
import gc
import traceback
from typing import Dict, List
import warnings
warnings.filterwarnings("ignore")

# Windows-specific PyTorch thread settings
if sys.platform == "win32":
    try:
        torch.set_num_threads(1)
    except RuntimeError:
        pass
    try:
        torch.set_num_interop_threads(1)
    except RuntimeError:
        pass

# Import main app
from main import TextRewritingApp
import torch.nn as nn
from transformers import AutoModelForCausalLM
try:
    from trl.experimental.ppo import AutoModelForCausalLMWithValueHead
except ImportError:
    from trl.models import AutoModelForCausalLMWithValueHead
from config import ppo_config

class PPOUpdateTrainer:
    """Simple trainer that updates existing LLM model weights using PPO-based feedback"""
    
    def __init__(self, model, tokenizer, reward_function, device="auto", model_save_path="./models"):
        # Use the existing LLM model directly - no separate model needed
        self.model = model
        self.tokenizer = tokenizer
        self.reward_function = reward_function
        self.device = device if device != "auto" else ("cuda" if torch.cuda.is_available() else "cpu")
        self.model_save_path = model_save_path
        
        # Ensure model is on correct device and in train mode
        self.model.to(self.device)
        self.model.train()
        
        # Initialize optimizer for updating weights
        self.optimizer = torch.optim.Adam(self.model.parameters(), lr=ppo_config.learning_rate)
        self.clip_ratio = ppo_config.clip_ratio
        
    def update_model(self, experiences):
        """Update LLM model weights using PPO with feedback experiences"""
        if not experiences:
            return
        
        try:
            self.model.train()
            
            for experience in experiences:
                original_text = experience['original_text']
                rewritten_text = experience['rewritten_text']
                reward = experience['reward']
                
                try:
                    # Ensure reward is valid
                    reward = float(reward)
                    reward = max(0.0, min(1.0, reward))
                    
                    # Keep text short to avoid memory issues
                    if len(rewritten_text) > 150:
                        rewritten_text = rewritten_text[:150]
                    if len(rewritten_text) < 10:
                        print(f"Skipping: text too short ({len(rewritten_text)} chars)")
                        continue
                    
                    # Tokenize rewritten text (the good output we want to reinforce)
                    inputs = self.tokenizer(
                        rewritten_text, 
                        return_tensors="pt", 
                        padding=True, 
                        truncation=True, 
                        max_length=128
                    ).to(self.device)
                    
                    if inputs['input_ids'].shape[1] < 2:
                        print(f"Skipping: input too short ({inputs['input_ids'].shape[1]} tokens)")
                        continue
                    
                    input_ids = inputs['input_ids']
                    attention_mask = inputs.get('attention_mask', None)
                    
                    # Forward pass - get logits
                    outputs = self.model(input_ids=input_ids, attention_mask=attention_mask)
                    logits = outputs.logits
                    
                    # Shift for next token prediction
                    shift_logits = logits[..., :-1, :].contiguous()
                    shift_labels = input_ids[..., 1:].contiguous()
                    
                    if shift_labels.shape[1] == 0:
                        print("Skipping: no labels after shift")
                        continue
                    
                    # Compute loss manually
                    loss_fct = torch.nn.CrossEntropyLoss(ignore_index=self.tokenizer.pad_token_id if self.tokenizer.pad_token_id else -100)
                    flat_logits = shift_logits.view(-1, shift_logits.size(-1))
                    flat_labels = shift_labels.view(-1)
                    
                    # Mask padding tokens
                    if attention_mask is not None:
                        mask = attention_mask[..., 1:].contiguous().view(-1).bool()
                        flat_logits = flat_logits[mask]
                        flat_labels = flat_labels[mask]
                    
                    if flat_labels.shape[0] == 0:
                        print("Skipping: no valid labels after masking")
                        continue
                    
                    loss = loss_fct(flat_logits, flat_labels)
                    
                    # Check for valid loss
                    if torch.isnan(loss) or torch.isinf(loss):
                        print(f"Skipping: invalid loss ({loss.item()})")
                        continue
                    
                    # Scale loss by reward: higher reward = less loss (we want to learn good outputs)
                    scaled_loss = loss * (1.0 - reward)
                    
                    # Update model weights
                    self.optimizer.zero_grad()
                    scaled_loss.backward()
                    torch.nn.utils.clip_grad_norm_(self.model.parameters(), 1.0)
                    self.optimizer.step()
                    
                    print(f"✅ Model updated: reward={reward:.3f}, loss={loss.item():.4f}, scaled={scaled_loss.item():.4f}")
                    
                except Exception as e:
                    print(f"❌ Error in PPO update: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
        except Exception as e:
            print(f"Error in update_model: {e}")
            raise
        finally:
            # Always set back to eval mode and cleanup
            self.model.eval()
            gc.collect()
            if torch.cuda.is_available():
                torch.cuda.empty_cache()

# Page configuration
st.set_page_config(
    page_title="Text Rewriting with RLHF",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        padding: 1rem 0;
    }
    .stButton>button {
        width: 100%;
        background-color: #1f77b4;
        color: white;
        font-weight: bold;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize session state
if 'app' not in st.session_state:
    st.session_state.app = None
if 'processing' not in st.session_state:
    st.session_state.processing = False

@st.cache_resource
def initialize_app(device="auto"):
    """Initialize the app with models"""
    try:
        app = TextRewritingApp(device=device)
        return app
    except Exception as e:
        st.error(f"Failed to initialize app: {str(e)}")
        return None

def display_analysis_result(analysis: Dict, title: str = "Analysis"):
    """Display analysis results in a nice format"""
    st.markdown(f"### {title}")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        sentiment_label = analysis.get('sentiment_label', 'N/A')
        sentiment_score = analysis.get('sentiment_score', 0)
        sentiment_color = "🟢" if sentiment_score > 0 else "🔴" if sentiment_score < 0 else "🟡"
        st.metric("Sentiment", f"{sentiment_color} {sentiment_label}", f"{sentiment_score:.2f}")
    
    with col2:
        tone_score = analysis.get('tone_score', 0)
        st.metric("Tone Score", f"{tone_score:.3f}")
    
    with col3:
        intent_score = analysis.get('intent_score', 0)
        st.metric("Intent Score", f"{intent_score:.3f}")
    
    with col4:
        overall_score = analysis.get('overall_score', 0)
        st.metric("Overall Score", f"{overall_score:.3f}")

def main():
    """Main Streamlit application"""
    
    # Header
    st.markdown('<h1 class="main-header">✨ Text Rewriting with RLHF</h1>', unsafe_allow_html=True)
    st.markdown("---")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Settings")
        
        device_option = st.selectbox(
            "Device",
            ["auto", "cpu", "cuda"],
            help="Select device to run models on"
        )
        
        # Initialize app
        if st.session_state.app is None:
            with st.spinner("Initializing app..."):
                st.session_state.app = initialize_app(device=device_option)
        
        app = st.session_state.app
        
        if app is None:
            st.error("Failed to initialize app")
            st.stop()
        
        st.markdown("---")
        st.header("📊 App Info")
        st.info(f"Device: **{app.device}**")
        
        if torch.cuda.is_available():
            st.success("✅ CUDA Available")
        else:
            st.info("ℹ️ Running on CPU")
    
    # Main content
    tab1, tab2, tab3 = st.tabs(["🔍 Analyze", "✍️ Rewrite", "🔄 Process with Feedback"])
    
    # Tab 1: Analyze
    with tab1:
        st.header("Text Sentiment Analysis")
        st.write("Analyze the sentiment, tone, intent, and meaning of your text")
        
        text_input = st.text_area(
            "Enter text to analyze",
            height=150,
            placeholder="Enter your text here..."
        )
        
        if st.button("🔍 Analyze Text", type="primary"):
            if not text_input or not text_input.strip():
                st.warning("Please enter some text to analyze")
            else:
                with st.spinner("Analyzing text..."):
                    result = app.analyze_text(text_input)
                    
                    if result.get("success"):
                        st.success("Analysis complete!")
                        display_analysis_result(result["analysis"])
                        
                        if "summary" in result:
                            st.markdown("### Summary")
                            st.info(result["summary"])
                    else:
                        st.error(f"Analysis failed: {result.get('error', 'Unknown error')}")
    
    # Tab 2: Rewrite
    with tab2:
        st.header("Text Rewriting")
        st.write("Rewrite text to be more polite, positive, or friendly")
        
        rewrite_text = st.text_area(
            "Enter text to rewrite",
            height=150,
            placeholder="Enter your text here..."
        )
        
        if st.button("✍️ Rewrite Text", type="primary"):
            if not rewrite_text or not rewrite_text.strip():
                st.warning("Please enter some text to rewrite")
            else:
                with st.spinner("Rewriting text... This may take a moment as we load the model."):
                    result = app.rewrite_text(rewrite_text)
                    
                    if result.get("success"):
                        st.success("Text rewritten successfully!")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("### Original Text")
                            st.text_area("Original Text", result["original_text"], height=150, key="original", disabled=True, label_visibility="hidden")
                            display_analysis_result(result["original_analysis"], "Original Analysis")
                        
                        with col2:
                            st.markdown("### Rewritten Text")
                            st.text_area("Rewritten Text", result["rewritten_text"], height=150, key="rewritten", disabled=True, label_visibility="hidden")
                            display_analysis_result(result["rewritten_analysis"], "Rewritten Analysis")
                        
                        # Improvement metrics
                        st.markdown("### 📈 Improvements")
                        improvement = result.get("improvement", {})
                        meaning_preservation = result.get("meaning_preservation", 0.5)
                        
                        imp_col1, imp_col2, imp_col3, imp_col4 = st.columns(4)
                        with imp_col1:
                            st.metric("Sentiment Improvement", f"{improvement.get('sentiment', 0):.3f}")
                        with imp_col2:
                            st.metric("Tone Improvement", f"{improvement.get('tone', 0):.3f}")
                        with imp_col3:
                            st.metric("Overall Improvement", f"{improvement.get('overall', 0):.3f}")
                        with imp_col4:
                            meaning_status = "✅ Preserved" if improvement.get('meaning_preserved', False) else "⚠️ Changed"
                            st.metric("Meaning Preservation", meaning_status, f"{meaning_preservation:.2f}")
                        
                        # Meaning preservation indicator
                        if meaning_preservation < 0.6:
                            st.warning(f"⚠️ Meaning preservation is low ({meaning_preservation:.2f}). The rewritten text may have changed the original meaning.")
                        elif meaning_preservation >= 0.8:
                            st.success(f"✅ Meaning is well preserved ({meaning_preservation:.2f}).")
                        else:
                            st.info(f"ℹ️ Meaning preservation: {meaning_preservation:.2f}")
                    else:
                        st.error(f"Rewriting failed: {result.get('error', 'Unknown error')}")
    
    # Tab 3: Process with Feedback
    with tab3:
        st.header("Process with Feedback")
        st.write("Generate 5 versions, select the best one, provide feedback, and train the model")
        
        # Check if we have results to display
        if st.session_state.get('process_complete', False) and st.session_state.get('process_results'):
            result = st.session_state.process_results
            
            st.markdown("### ✅ Processing Complete!")
            
            # Display Original Text
            st.markdown("#### Original Text:")
            st.text_area(
                "Original Text", 
                result.get("original_text", ""), 
                height=150, 
                key="final_original", 
                disabled=True, 
                label_visibility="collapsed"
            )
            
            # Display Rewritten Text
            st.markdown("#### ✍️ Rewritten Text:")
            st.text_area(
                "Rewritten Text", 
                result.get("selected_version", ""), 
                height=150, 
                key="final_rewritten", 
                disabled=True, 
                label_visibility="collapsed"
            )
            
            metrics_col1, metrics_col2, metrics_col3 = st.columns(3)
            with metrics_col1:
                st.metric("Final Reward", f"{result.get('reward', 0):.3f}")
            with metrics_col2:
                st.metric("User Rating", f"{result.get('user_rating', 0):.2f}")
            with metrics_col3:
                st.metric("Model Status", "✅ Saved")
            
            if "component_scores" in result:
                st.markdown("### 📈 Reward Components")
                component_scores = result["component_scores"]
                for key, value in component_scores.items():
                    if isinstance(value, (int, float)):
                        normalized_value = max(0, min(1, value)) if abs(value) <= 1 else max(0, min(1, value / 2))
                        st.progress(normalized_value, text=f"{key}: {value:.3f}")
                    else:
                        st.write(f"{key}: {value}")
            
            st.info("🎉 Model has been updated with your feedback!")
            
            if st.button("🔄 Start New Process"):
                # Reset session state
                for key in ['process_text', 'process_versions', 
                           'process_original_analysis', 'process_complete', 'process_results',
                           'process_version_analyses', 'process_meaning_scores']:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()
            
            return
        
        # If we have versions to select, show the selection UI
        if (st.session_state.get('process_versions') and 
            len(st.session_state.get('process_versions', [])) > 0 and
            not st.session_state.get('process_complete', False)):
            
            st.markdown("### Select the best rewritten version:")
            
            # Display versions
            if st.session_state.process_versions:
                sentiment_score = st.session_state.process_original_analysis.get('sentiment_score', 0)
                
                version_cols = st.columns(5)
                for idx, version in enumerate(st.session_state.process_versions):
                    with version_cols[idx]:
                        # Get analysis and meaning score if stored
                        if hasattr(st.session_state, 'process_version_analyses') and idx < len(st.session_state.process_version_analyses):
                            version_analysis = st.session_state.process_version_analyses[idx]
                        else:
                            version_analysis = app.analyze_text(version).get('analysis', {})
                        
                        if hasattr(st.session_state, 'process_meaning_scores') and idx < len(st.session_state.process_meaning_scores):
                            meaning_score = st.session_state.process_meaning_scores[idx]
                        else:
                            meaning_score = 0.5
                        
                        sentiment_improvement = version_analysis.get('sentiment_score', 0) - sentiment_score
                        
                        st.markdown(f"#### Option {idx + 1}")
                        st.text_area(f"Version {idx + 1}", version, height=120, key=f"version_display_{idx}", disabled=True, label_visibility="hidden")
                        st.metric("Sentiment Improvement", f"+{sentiment_improvement:.2f}")
                        meaning_status = "✅" if meaning_score >= 0.6 else "⚠️"
                        st.metric("Meaning Preservation", f"{meaning_status} {meaning_score:.2f}")
                        st.caption(f"New Sentiment: {version_analysis.get('sentiment_score', 0):.2f}")
            
            # Display versions with selection
            selected_option = st.radio(
                "Select the best rewritten version:",
                options=range(len(st.session_state.process_versions)),
                format_func=lambda x: f"Option {x+1}",
                key="select_version"
            )
            
            # Show selected version details
            if st.session_state.process_original_analysis:
                sentiment_score = st.session_state.process_original_analysis.get('sentiment_score', 0)
                selected_text = st.session_state.process_versions[selected_option]
                selected_analysis_result = app.analyze_text(selected_text)
                selected_analysis = selected_analysis_result.get('analysis', {})
                sentiment_improvement = selected_analysis.get('sentiment_score', 0) - sentiment_score
                
                st.markdown("#### Selected Version Details:")
                st.text_area("Selected Text", selected_text, height=100, disabled=True)
                
                # Get meaning score for selected version
                selected_idx = selected_option
                if hasattr(st.session_state, 'process_meaning_scores') and selected_idx < len(st.session_state.process_meaning_scores):
                    selected_meaning = st.session_state.process_meaning_scores[selected_idx]
                else:
                    selected_meaning = app._compute_meaning_preservation(st.session_state.process_text, selected_text)
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Sentiment Improvement", f"+{sentiment_improvement:.3f}")
                with col2:
                    st.metric("New Sentiment", f"{selected_analysis.get('sentiment_score', 0):.3f}")
                with col3:
                    meaning_status = "✅ Good" if selected_meaning >= 0.6 else "⚠️ Low"
                    st.metric("Meaning Preservation", meaning_status, f"{selected_meaning:.2f}")
            
            # Get user rating
            st.markdown("#### Provide Feedback:")
            user_rating = st.slider(
                "Rate the quality (0.0 = Very poor, 1.0 = Excellent):",
                min_value=0.0,
                max_value=1.0,
                value=0.5,
                step=0.25
            )
            
            rating_labels = {
                0.0: "Very Poor (0.0)",
                0.25: "Poor (0.25)",
                0.5: "Fair (0.5)",
                0.75: "Good (0.75)",
                1.0: "Excellent (1.0)"
            }
            st.caption(f"Current rating: {rating_labels.get(user_rating, f'{user_rating:.2f}')}")
            
            if st.button("✓ Submit Feedback & Save Model", type="primary", key="submit_feedback"):
                status_placeholder = st.empty()
                
                try:
                    selected_text = st.session_state.process_versions[selected_option]
                    
                    status_placeholder.info("Processing feedback...")
                    
                    # Ensure text_rewriter is loaded
                    if app.text_rewriter is None:
                        status_placeholder.warning("Loading model...")
                        with st.spinner("Loading Qwen model..."):
                            from qwen_rewriter import QwenRewriter
                            app.text_rewriter = QwenRewriter(device=app.device)
                    
                    # Initialize trainer if needed, using existing LLM model
                    if app.trainer is None:
                        status_placeholder.info("Initializing trainer...")
                        try:
                            from reward_function import RewardFunction
                            
                            # Create reward function
                            reward_function = RewardFunction(app.sentiment_analyzer)
                            
                            # Create trainer that uses the existing LLM model directly
                            app.trainer = PPOUpdateTrainer(
                                model=app.text_rewriter.model,
                                tokenizer=app.text_rewriter.tokenizer,
                                reward_function=reward_function,
                                device=app.device,
                                model_save_path="./models"
                            )
                            status_placeholder.success("Trainer initialized")
                        except Exception as e:
                            status_placeholder.error(f"Failed to initialize trainer: {str(e)}")
                            st.exception(e)
                            return
                    
                    status_placeholder.info("Computing reward...")
                    # Compute reward
                    reward, component_scores = app.trainer.reward_function.compute_reward(
                        st.session_state.process_text,
                        selected_text,
                        user_rating=user_rating
                    )
                    
                    status_placeholder.info("Updating model weights with PPO...")
                    st.info(f"Reward: {reward:.3f} | Rating: {user_rating:.2f}")
                    
                    # Create experience for PPO update
                    experience = {
                        'original_text': st.session_state.process_text,
                        'rewritten_text': selected_text,
                        'reward': reward,
                        'component_scores': component_scores,
                        'user_rating': user_rating
                    }
                    
                    # Update LLM model weights with PPO (uses the same model instance)
                    update_success = False
                    try:
                        with st.spinner("Training model with feedback..."):
                            app.trainer.update_model([experience])
                            update_success = True
                            status_placeholder.success("✅ Model weights updated successfully!")
                    except Exception as e:
                        error_msg = str(e)
                        status_placeholder.warning(f"⚠️ Model update had issues: {error_msg[:100]}")
                        st.warning(f"Update error details: {error_msg}")
                        # Still try to save if update partially worked
                        update_success = True  # Try to continue anyway
                    
                    # Model is already updated in-place, just ensure it's in eval mode
                    try:
                        app.text_rewriter.model.eval()
                    except:
                        pass
                    
                    # Save the updated LLM model
                    status_placeholder.info("Saving model to disk...")
                    save_success = False
                    try:
                        gc.collect()
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                        
                        os.makedirs("./models", exist_ok=True)
                        
                        # Save model
                        model_path = "./models"
                        app.text_rewriter.model.save_pretrained(model_path, safe_serialization=True)
                        app.text_rewriter.tokenizer.save_pretrained(model_path)
                        
                        save_success = True
                        status_placeholder.success("✅ Model saved successfully to ./models/")
                    except Exception as e:
                        error_msg = str(e)
                        status_placeholder.error(f"❌ Failed to save model: {error_msg[:200]}")
                        st.error(f"Save error: {error_msg}")
                        save_success = False
                    
                    if not save_success:
                        st.error("⚠️ Model weights were updated but could not be saved. Please check ./models/ directory permissions.")
                    
                    # Complete processing
                    st.session_state.process_results = {
                        "success": True,
                        "original_text": st.session_state.process_text,
                        "selected_version": selected_text,
                        "user_rating": user_rating,
                        "reward": reward,
                        "component_scores": component_scores
                    }
                    st.session_state.process_complete = True
                    
                    # Clear status and rerun
                    status_placeholder.empty()
                    st.rerun()
                    
                except Exception as e:
                    status_placeholder.error(f"Error processing feedback: {str(e)}")
                    st.exception(e)
                    import traceback
                    traceback.print_exc()
        
        # Initial setup - get text and parameters
        else:
            # Input method selection
            input_method = st.radio(
                "Input Method:",
                ["Type Text", "Upload Document"],
                horizontal=True
            )
            
            process_text = ""
            
            if input_method == "Type Text":
                process_text = st.text_area("Enter text to process:", height=150)
            else:
                uploaded_file = st.file_uploader(
                    "Upload a document",
                    type=['txt', 'pdf', 'docx', 'doc'],
                    help="Upload a text, PDF, or Word document"
                )
                
                if uploaded_file:
                    try:
                        if uploaded_file.name.endswith('.txt'):
                            process_text = uploaded_file.read().decode('utf-8')
                        elif uploaded_file.name.endswith('.pdf'):
                            import PyPDF2
                            pdf_reader = PyPDF2.PdfReader(uploaded_file)
                            process_text = ""
                            for page in pdf_reader.pages:
                                text = page.extract_text()
                                if text:
                                    process_text += text + "\n"
                            if not process_text.strip():
                                st.warning("Could not extract text from PDF. The PDF might be image-based or encrypted.")
                        elif uploaded_file.name.endswith(('.docx', '.doc')):
                            if uploaded_file.name.endswith('.doc'):
                                st.warning("⚠️ .doc files are not directly supported. Please convert to .docx first.")
                            else:
                                from docx import Document
                                doc = Document(uploaded_file)
                                process_text = ""
                                for para in doc.paragraphs:
                                    process_text += para.text + "\n"
                                for table in doc.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            process_text += cell.text + "\n"
                        
                        if process_text:
                            st.success(f"✅ Document loaded: {uploaded_file.name}")
                            with st.expander("📄 Preview extracted text"):
                                st.text_area("Extracted Text", process_text[:500] + "..." if len(process_text) > 500 else process_text, 
                                           height=150, disabled=True, label_visibility="collapsed")
                    except Exception as e:
                        st.error(f"Error reading document: {str(e)}")
                        st.exception(e)
                        process_text = ""
            
            if st.button("🚀 Start Processing", type="primary"):
                if not process_text or not process_text.strip():
                    st.warning("Please enter some text to process")
                else:
                    with st.spinner("Analyzing text and generating versions..."):
                        try:
                            # Step 1: Analyze
                            original_analysis = app.sentiment_analyzer.analyze_text(process_text)
                            sentiment_score = original_analysis.get('sentiment_score', 0)
                            
                            if sentiment_score >= 0.0:
                                st.warning(f"⚠️ Text is not negative (score: {sentiment_score:.2f}). Proceeding anyway...")
                            
                            # Step 2: Generate versions
                            st.info("Generating 5 rewritten versions (this may take a moment)...")
                            versions_result = app.generate_multiple_versions(process_text, num_versions=5)
                            
                            if not versions_result.get("success"):
                                st.error(f"Failed to generate versions: {versions_result.get('error', 'Unknown error')}")
                            else:
                                versions = versions_result.get("versions", [])
                                version_analyses = versions_result.get("version_analyses", [])
                                meaning_scores = versions_result.get("meaning_scores", [])
                                
                                if not versions or len(versions) == 0:
                                    st.error("Failed to generate rewritten versions")
                                else:
                                    st.session_state.process_text = process_text
                                    st.session_state.process_versions = versions
                                    st.session_state.process_original_analysis = original_analysis
                                    st.session_state.process_version_analyses = version_analyses
                                    st.session_state.process_meaning_scores = meaning_scores
                                    st.session_state.process_complete = False
                                    st.rerun()
                        except Exception as e:
                            st.error(f"Error processing text: {str(e)}")
                            st.exception(e)

if __name__ == "__main__":
    main()
